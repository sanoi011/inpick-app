"""
INPICK 웹/앱 결제 대표님 작업 가이드 — .docx 생성.

생성 위치: C:\\Users\\user\\Desktop\\INPICK_결제_대표님_작업_가이드_20260519.docx
"""
from __future__ import annotations
import io, os, sys
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[install] python-docx 설치 필요: pip install python-docx")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ─── 스타일 설정 ─────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10.5)
    if bold:
        run.bold = True
    return p

def bullet(text, indent=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
    return p

def numbered(text, indent=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    return t

# ═════════════════════════════════════════════════
# 표지
# ═════════════════════════════════════════════════
title = doc.add_heading("INPICK 웹/앱 결제 대표님 작업 가이드", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("2026-05-19 작성 · 외부 콘솔 작업 + 환경변수 등록 순서")
r.font.size = Pt(11)
r.italic = True
doc.add_paragraph()

body(
    "이 문서는 Claude Code(저)가 만들 수 없고 대표님이 직접 외부 콘솔에서 해야 하는 작업만 "
    "단계별로 정리한 가이드입니다. 코드/DB/API는 모두 작성 완료되었으며 환경변수만 등록하면 "
    "즉시 동작합니다."
)
body(
    "체크박스 형식으로 진행하시면서 막히는 부분 있으면 알려주세요."
)
doc.add_page_break()

# ═════════════════════════════════════════════════
# 0. 한눈에 보는 작업 순서
# ═════════════════════════════════════════════════
heading("0. 한눈에 보는 작업 순서 (우선순위 순)", level=1)

body("아래 순서대로 진행하세요. 앞 단계 미완료 시 다음 단계 의미 없음.", bold=True)
add_table(
    ["순서", "단계", "예상 소요", "비용", "필수여부"],
    [
        ["1", "Toss 입점 통과 후 환경변수 등록 (웹 결제)", "1~2일", "0원", "P0 필수"],
        ["2", "Resend SMTP 등록 (비번 찾기 메일)", "30분", "월 3,000건 무료", "P1 권장"],
        ["3", "Apple Developer Program 가입", "1일", "$99/년", "iOS 출시 필수"],
        ["4", "App Store Connect 앱 등록 + Bundle ID", "1~2시간", "0원", "iOS 출시 필수"],
        ["5", "App Store Server API Key 발급", "30분", "0원", "iOS 결제 필수"],
        ["6", "IAP 상품 등록 (App Store Connect)", "1시간", "0원", "iOS 결제 필수"],
        ["7", "Google Play Console 개발자 가입", "1~2일 (심사)", "$25 일회성", "Android 출시 필수"],
        ["8", "Play Console 앱 등록 + 패키지명", "1~2시간", "0원", "Android 출시 필수"],
        ["9", "Google Play Service Account 발급", "30분", "0원", "Android 결제 필수"],
        ["10", "Play Billing 상품 등록", "1시간", "0원", "Android 결제 필수"],
        ["11", "Vercel 환경변수 등록 (위 키 모두)", "30분", "0원", "최종 활성화"],
    ],
)
doc.add_page_break()

# ═════════════════════════════════════════════════
# 1. 웹 결제 — 토스페이먼츠
# ═════════════════════════════════════════════════
heading("1. 웹 결제 — 토스페이먼츠 (가장 시급)", level=1)

heading("1-1. 현재 상태", level=2)
bullet("토스페이먼츠 입점 심사 중 (대표님 진행)")
bullet("코드는 모두 작성 완료 — Mock 모드로 동작 중")
bullet("입점 통과 + 키 발급 후 환경변수만 등록하면 즉시 실결제 전환")

heading("1-2. 토스 입점 통과 후 받아야 할 것", level=2)
bullet("TOSS_PAYMENTS_CLIENT_KEY (live_ck_xxxx... 형식)")
bullet("TOSS_PAYMENTS_SECRET_KEY (live_sk_xxxx... 형식)")
bullet("TOSS_PAYMENTS_WEBHOOK_SECRET (선택, 웹훅 서명 검증용)")
bullet("토스 가맹점 식별자 (MID)")

heading("1-3. Vercel 환경변수 등록 방법", level=2)
numbered("https://vercel.com/sanoi011s-projects/inpick-app/settings/environment-variables 접속")
numbered("Add New 클릭")
numbered("아래 키들을 Production 환경에 등록:")
code_block(
    "TOSS_PAYMENTS_CLIENT_KEY = live_ck_xxxx\n"
    "TOSS_PAYMENTS_SECRET_KEY = live_sk_xxxx\n"
    "NEXT_PUBLIC_TOSS_CLIENT_KEY = live_ck_xxxx  (클라이언트 노출용 — 동일 값)\n"
    "TOSS_WEBHOOK_SECRET = (받은 webhook secret)"
)
numbered("Save 후 자동 재배포 (Vercel이 자동으로 main branch 재빌드)")
numbered("/admin/settings 페이지 접속해서 'Mock' → '라이브' 표시 확인")

body("⚠️ 주의:", bold=True)
bullet("test_ck_ / test_sk_ 키는 테스트용입니다. 운영용은 live_ck_ / live_sk_ 입니다.")
bullet("키를 절대 GitHub에 commit하지 마세요. Vercel UI에서만 등록.")

heading("1-4. 토스 결제 동작 확인", level=2)
numbered("환경변수 등록 후 https://inpick-app.vercel.app/mypage/billing 접속")
numbered("'토큰 충전' 버튼 클릭 → 패키지 선택 → 결제")
numbered("실제 카드/계좌 결제창 표시되면 정상")
numbered("결제 완료 후 /admin/payment-center 에서 결제 내역 확인")
doc.add_page_break()

# ═════════════════════════════════════════════════
# 2. Resend SMTP — 메일 발송
# ═════════════════════════════════════════════════
heading("2. Resend SMTP — 비밀번호 찾기 메일 (권장)", level=1)

heading("2-1. 왜 필요한가", level=2)
body(
    "현재 Supabase 기본 SMTP는 시간당 4건 제한 + 신뢰성 낮음. 비밀번호 찾기 메일이 미수신되는 "
    "이유가 이것입니다. (어제 자체 비번 재설정 흐름을 만들었으므로 SMTP 없이도 복구 가능하지만, "
    "메일 인증/알림용으로는 SMTP가 필요합니다.)"
)

heading("2-2. Resend 가입 + 도메인 인증", level=2)
numbered("https://resend.com 가입 (Google/GitHub 로그인 가능, 카드 등록 불필요)")
numbered("Domains → Add Domain → inpick.kr (또는 별도 도메인) 입력")
numbered("DNS 레코드 4개를 도메인 등록업체(가비아 등)에 추가:")
code_block(
    "TXT _resend.inpick.kr → resend-domain-verification=xxx\n"
    "MX send.inpick.kr → feedback-smtp.us-east-1.amazonses.com\n"
    "TXT send.inpick.kr → v=spf1 include:amazonses.com ~all\n"
    "TXT resend._domainkey.inpick.kr → (DKIM key — Resend가 제공)"
)
numbered("DNS 전파 후 Resend가 'Verified' 표시")
body("⚠️ 도메인 등록 안 된 상태라면 임시로 onboarding@resend.dev 발신 가능 (월 100건)")

heading("2-3. API Key 발급", level=2)
numbered("Resend → API Keys → Create API Key")
numbered("이름: inpick-production-smtp")
numbered("권한: Send emails")
numbered("발급된 키 (re_xxxxx) 복사")

heading("2-4. Supabase SMTP 설정", level=2)
numbered("https://supabase.com/dashboard/project/pyhsjjtxcfmkcqmaxozd/settings/auth 접속")
numbered("'Enable Custom SMTP' 토글 ON")
numbered("아래 입력:")
code_block(
    "Host: smtp.resend.com\n"
    "Port: 587 (또는 465 SSL)\n"
    "User: resend\n"
    "Password: re_xxxxx (위에서 발급한 API Key)\n"
    "Sender email: noreply@inpick.kr (인증된 도메인)\n"
    "Sender name: INPICK"
)
numbered("Save Changes")

heading("2-5. 메일 템플릿 한국어화 (선택)", level=2)
body(
    "Supabase → Authentication → Email Templates 에서 6개 템플릿 한국어로 수정 권장. "
    "기본은 영어이며, 회원가입 확인/비번 재설정/이메일 변경 등 자동 발송 시 사용됩니다."
)
doc.add_page_break()

# ═════════════════════════════════════════════════
# 3. Apple App Store (iOS 앱)
# ═════════════════════════════════════════════════
heading("3. Apple App Store — iOS 앱 출시 + 결제", level=1)

heading("3-1. Apple Developer Program 가입", level=2)
numbered("https://developer.apple.com/programs/enroll 접속")
numbered("Apple ID 로그인 (없으면 생성)")
numbered("개인 또는 법인 선택:")
bullet("법인 (대영토건): D-U-N-S Number 필요 (https://developer.apple.com/support/D-U-N-S/ 신청, 무료, 5~10일)", indent=1)
bullet("개인: 즉시 가능, 단 앱 이름에 회사명 표시 제한", indent=1)
numbered("$99/년 결제 (자동 갱신, 카드 필요)")
numbered("가입 완료까지 1~2일 소요")

heading("3-2. App Store Connect 앱 등록", level=2)
numbered("https://appstoreconnect.apple.com 접속")
numbered("My Apps → + → New App")
numbered("아래 입력:")
code_block(
    "Platform: iOS\n"
    "Name: INPICK (또는 인픽)\n"
    "Primary Language: Korean\n"
    "Bundle ID: kr.inpick.app  ← 반드시 이 ID로\n"
    "SKU: inpick-ios-001  (내부 식별, 자유)\n"
    "User Access: Full Access"
)
numbered("Create")

heading("3-3. Bundle ID 사전 등록 (Bundle ID가 없다고 나오면)", level=2)
numbered("https://developer.apple.com/account/resources/identifiers 접속")
numbered("App IDs → + → App")
numbered("Description: INPICK Consumer App")
numbered("Bundle ID: Explicit → kr.inpick.app")
numbered("Capabilities: Sign In with Apple, In-App Purchase 체크")
numbered("Continue → Register")

heading("3-4. App Store Server API Key 발급 (결제 검증용)", level=2)
numbered("https://appstoreconnect.apple.com/access/api 접속")
numbered("Integrations 탭 → Generate API Key")
numbered("Name: inpick-iap-verify")
numbered("Access: App Manager")
numbered("Download .p8 파일 (한 번만 다운로드 가능, 분실 시 재발급 X)")
numbered("Key ID + Issuer ID 메모")

body("Vercel에 등록할 환경변수:", bold=True)
code_block(
    "APP_STORE_BUNDLE_ID = kr.inpick.app\n"
    "APP_STORE_KEY_ID = (위 Key ID)\n"
    "APP_STORE_ISSUER_ID = (위 Issuer ID)\n"
    "APP_STORE_PRIVATE_KEY = (다운로드한 .p8 파일 내용을 base64 인코딩한 값)\n"
    "APP_STORE_ENV = sandbox  (테스트) → production  (출시 후)"
)
body(
    "💡 .p8 파일을 base64로 인코딩하는 방법 (Windows PowerShell): "
    "[Convert]::ToBase64String([IO.File]::ReadAllBytes('AuthKey_XXXX.p8'))"
)

heading("3-5. IAP (In-App Purchase) 상품 등록", level=2)
body("App Store Connect → My Apps → INPICK → In-App Purchases → +")
body("아래 4개 토큰 패키지를 'Consumable'로 등록:")
add_table(
    ["Product ID", "Reference Name", "가격 (Tier)", "지급 토큰"],
    [
        ["kr.inpick.token.10", "AI Token 10", "Tier 5 (₩5,000)", "10"],
        ["kr.inpick.token.33", "AI Token 33 (30+3)", "Tier 15 (₩15,000)", "33"],
        ["kr.inpick.token.115", "AI Token 115 (100+15)", "Tier 50 (₩50,000)", "115"],
        ["kr.inpick.token.360", "AI Token 360 (300+60)", "Tier 150 (₩150,000)", "360"],
    ],
)
body("PDF 발급권 (Non-Consumable):")
code_block(
    "Product ID: kr.inpick.pdf.single\n"
    "Type: Non-Consumable\n"
    "Tier: Tier 9 (₩9,900)"
)
body("각 상품마다 한국어/영어 표시 이름 + 설명 작성. 심사 통과까지 1~2일.")

heading("3-6. TestFlight 베타 배포 (앱 1차 출시 전 권장)", level=2)
numbered("Xcode에서 앱 archive → Upload to App Store Connect")
numbered("App Store Connect → TestFlight → 빌드 자동 등록 대기 (5~10분)")
numbered("Internal Testing 그룹 생성 + 본인 Apple ID 추가")
numbered("TestFlight 앱 (iOS)으로 본인 폰에서 설치 + 테스트")
numbered("외부 테스터는 별도 그룹 + Apple 심사 필요 (24시간)")

heading("3-7. 앱 심사 제출 (1차 출시)", level=2)
body("App Store Connect → App Information + Pricing + App Review Information 입력:")
bullet("개인정보처리방침 URL: https://inpick.kr/privacy (또는 vercel URL)")
bullet("심사 정보 — 데모 계정 (심사관용 로그인 계정 필수)")
bullet("연락처 (이메일/전화)")
bullet("앱 카테고리: 라이프스타일 또는 비즈니스")
bullet("연령 등급: 4+ (별도 폭력/성적 콘텐츠 없으면)")
bullet("스크린샷 4개 이상 (iPhone 6.7\", 6.5\", 5.5\" 각각)")
bullet("앱 아이콘 (1024x1024 PNG)")
bullet("Submit for Review → 심사 1~7일 소요")
doc.add_page_break()

# ═════════════════════════════════════════════════
# 4. Google Play (Android 앱)
# ═════════════════════════════════════════════════
heading("4. Google Play — Android 앱 출시 + 결제", level=1)

heading("4-1. Google Play Console 개발자 등록", level=2)
numbered("https://play.google.com/console 접속")
numbered("Google 계정 로그인")
numbered("'개발자 계정 만들기' (개인 또는 조직)")
bullet("개인: 즉시 가능, 본인 인증 필요", indent=1)
bullet("조직: D-U-N-S 또는 사업자등록증, 1~2주 심사", indent=1)
numbered("$25 일회성 결제 (평생, 카드 필요)")
numbered("가입 + 본인 인증까지 1~2일")

heading("4-2. 앱 등록", level=2)
numbered("Play Console → Create app")
numbered("아래 입력:")
code_block(
    "앱 이름: INPICK\n"
    "기본 언어: 한국어\n"
    "앱/게임: 앱\n"
    "유료/무료: 무료\n"
    "패키지 이름: kr.inpick.app  ← 반드시 동일\n"
    "선언: 모두 체크"
)

heading("4-3. Play App Signing (필수)", level=2)
body(
    "Play App Signing은 Google이 앱 서명 키를 관리하는 기능. 키 분실 위험 없고 권장 방식."
)
numbered("Play Console → Setup → App integrity → App signing")
numbered("'Use Play App Signing' 선택")
numbered("Upload key 생성 (Android Studio 또는 keytool):")
code_block(
    "keytool -genkey -v -keystore inpick-upload.keystore \\\n"
    "  -alias inpick-upload -keyalg RSA -keysize 2048 -validity 10000"
)
numbered("inpick-upload.keystore 파일 안전한 곳에 백업 (분실 시 앱 업데이트 불가)")
numbered("Upload key의 SHA-1 fingerprint를 Play Console에 등록")

heading("4-4. Google Play Service Account 발급 (결제 검증용)", level=2)
numbered("https://console.cloud.google.com 접속 (Google Play Console과 다른 사이트)")
numbered("프로젝트 생성: 'inpick-play-billing' (또는 기존 사용)")
numbered("APIs & Services → Library → 'Google Play Android Developer API' 검색 → Enable")
numbered("IAM & Admin → Service Accounts → Create Service Account")
numbered("이름: inpick-iap-verify, 역할: 없음 (Play Console에서 별도 권한 부여)")
numbered("Keys → Add Key → JSON 다운로드 (이 JSON 파일은 한 번만 받을 수 있음)")
numbered("다시 Play Console로 돌아가서:")
bullet("Play Console → Setup → API access → Service accounts에 위 SA 이메일 추가", indent=1)
bullet("권한: View financial data, Manage orders and subscriptions 부여", indent=1)
numbered("권한 전파까지 24시간 소요 가능")

body("Vercel 환경변수:", bold=True)
code_block(
    "GOOGLE_PLAY_PACKAGE_NAME = kr.inpick.app\n"
    "GOOGLE_PLAY_SERVICE_ACCOUNT_JSON = (위 JSON 파일 전체를 base64 인코딩한 값)"
)
body(
    "💡 JSON을 base64로 인코딩 (PowerShell): "
    "[Convert]::ToBase64String([Text.Encoding]::UTF8.GetBytes((Get-Content service-account.json -Raw)))"
)

heading("4-5. Play Billing 상품 등록", level=2)
body("Play Console → Monetize → Products → In-app products → Create product")
add_table(
    ["Product ID", "이름", "가격", "지급 토큰"],
    [
        ["kr.inpick.token.10", "AI 토큰 10개", "₩5,000", "10"],
        ["kr.inpick.token.33", "AI 토큰 33개", "₩15,000", "33"],
        ["kr.inpick.token.115", "AI 토큰 115개", "₩50,000", "115"],
        ["kr.inpick.token.360", "AI 토큰 360개", "₩150,000", "360"],
        ["kr.inpick.pdf.single", "견적서 PDF 발급권", "₩9,900", "1회권"],
    ],
)
body("각 상품 'Activate' 클릭 + 출시 후에야 실제 구매 가능")

heading("4-6. 내부 테스트 트랙", level=2)
numbered("Play Console → Testing → Internal testing → Create release")
numbered("AAB(Android App Bundle) 업로드 (Android Studio에서 build)")
numbered("Testers 추가 (Google 그룹 또는 이메일 직접)")
numbered("Test URL 받아 본인 Android 폰에서 설치")

heading("4-7. 프로덕션 출시 심사 제출", level=2)
bullet("App content 섹션 모두 작성:")
bullet("개인정보처리방침 URL", indent=1)
bullet("Data Safety (수집 데이터 종류, 공유 여부, 사용 목적)", indent=1)
bullet("콘텐츠 등급 설문", indent=1)
bullet("타겟 사용자: 만 18세 이상", indent=1)
bullet("광고: 없음", indent=1)
bullet("Store listing — 스크린샷, 아이콘, 설명")
bullet("Pricing & distribution — 한국 + 국가 선택")
bullet("Production release → Submit → 심사 3~7일")
doc.add_page_break()

# ═════════════════════════════════════════════════
# 5. 환경변수 등록 최종 체크리스트
# ═════════════════════════════════════════════════
heading("5. Vercel 환경변수 최종 체크리스트", level=1)

body("아래 모든 키를 Production 환경에 등록 (Vercel → Settings → Environment Variables):")

heading("5-1. 결제 (필수)", level=2)
add_table(
    ["환경변수", "값 형식", "용도"],
    [
        ["TOSS_PAYMENTS_CLIENT_KEY", "live_ck_xxxx", "Toss 웹 결제"],
        ["TOSS_PAYMENTS_SECRET_KEY", "live_sk_xxxx", "Toss 서버 confirm"],
        ["NEXT_PUBLIC_TOSS_CLIENT_KEY", "live_ck_xxxx", "Toss 클라이언트 SDK"],
        ["TOSS_WEBHOOK_SECRET", "(webhook secret)", "Toss webhook 서명 검증"],
    ],
)

heading("5-2. iOS 결제", level=2)
add_table(
    ["환경변수", "값", "비고"],
    [
        ["APP_STORE_BUNDLE_ID", "kr.inpick.app", "고정"],
        ["APP_STORE_KEY_ID", "10자리 영숫자", "App Store Connect API Key ID"],
        ["APP_STORE_ISSUER_ID", "UUID 형식", "Issuer ID"],
        ["APP_STORE_PRIVATE_KEY", "base64(.p8 파일)", "다운로드 한 번만 가능"],
        ["APP_STORE_ENV", "sandbox 또는 production", "테스트 → 운영"],
    ],
)

heading("5-3. Android 결제", level=2)
add_table(
    ["환경변수", "값", "비고"],
    [
        ["GOOGLE_PLAY_PACKAGE_NAME", "kr.inpick.app", "고정"],
        ["GOOGLE_PLAY_SERVICE_ACCOUNT_JSON", "base64(SA JSON)", "Service Account credential"],
    ],
)

heading("5-4. 메일 SMTP (Resend)", level=2)
body("Supabase 대시보드에서 직접 설정 — Vercel 환경변수 X")
doc.add_page_break()

# ═════════════════════════════════════════════════
# 6. 작업 완료 체크리스트
# ═════════════════════════════════════════════════
heading("6. 단계별 진행 체크리스트", level=1)

heading("Phase 1 — 웹 결제 활성화 (이번 주)", level=2)
bullet("☐ Toss 입점 통과 확인")
bullet("☐ Vercel에 TOSS_* 환경변수 4개 등록")
bullet("☐ /admin/settings 에서 'Live' 모드 확인")
bullet("☐ 본인 계정으로 실제 결제 테스트 (5,000원 토큰 1패키지)")
bullet("☐ /admin/payment-center에서 결제 내역 확인")

heading("Phase 2 — 메일 SMTP (이번 주)", level=2)
bullet("☐ Resend 가입 + 도메인 인증 (또는 onboarding@resend.dev 임시)")
bullet("☐ Resend API Key 발급")
bullet("☐ Supabase Custom SMTP 설정")
bullet("☐ 본인 이메일로 비번 재설정 메일 수신 테스트")

heading("Phase 3 — iOS 앱 출시 준비 (1~2개월)", level=2)
bullet("☐ Apple Developer Program 가입 + 결제")
bullet("☐ App Store Connect 앱 등록 (Bundle ID: kr.inpick.app)")
bullet("☐ App Store Server API Key 발급 + .p8 다운로드 + 백업")
bullet("☐ IAP 상품 5개 등록 (토큰 4 + PDF 1)")
bullet("☐ Vercel APP_STORE_* 환경변수 5개 등록")
bullet("☐ TestFlight 내부 테스트")
bullet("☐ App Store 심사 제출")

heading("Phase 4 — Android 앱 출시 준비 (1~2개월)", level=2)
bullet("☐ Google Play Console 가입 + 결제")
bullet("☐ 앱 등록 (패키지명: kr.inpick.app)")
bullet("☐ Play App Signing 설정 + upload keystore 백업")
bullet("☐ Service Account 발급 + JSON 다운로드 + 백업")
bullet("☐ Play Billing 상품 5개 등록")
bullet("☐ Vercel GOOGLE_PLAY_* 환경변수 2개 등록")
bullet("☐ Internal testing 트랙 빌드 업로드")
bullet("☐ Production 심사 제출")
doc.add_page_break()

# ═════════════════════════════════════════════════
# 7. 자주 묻는 질문
# ═════════════════════════════════════════════════
heading("7. FAQ — 자주 막히는 부분", level=1)

heading("Q1. Apple Developer 법인 vs 개인 어느 게 좋나?", level=3)
body(
    "법인 (대영토건) 등록을 권장합니다. 이유: 앱 이름에 회사명 표시 가능, 세금 처리 명확, "
    "추후 사업자 변경 시 안전. 단, D-U-N-S 신청에 5~10일 소요됨. 시간 급하면 개인으로 시작 후 "
    "추후 법인으로 이관도 가능."
)

heading("Q2. Google Play 개인 등록 시 본명 노출되나?", level=3)
body(
    "예, 개인 등록 시 개발자 이름이 사용자에게 보입니다. 법인 등록 시 회사명만 보입니다. "
    "프라이버시 우려되면 법인 우선 권장."
)

heading("Q3. 앱 출시 안 하고 코드만 미리 만들어도 되나?", level=3)
body(
    "예, 가능합니다. 현재 모든 결제 검증 API + DB 인프라가 완료되어 있으므로 환경변수만 등록하면 "
    "즉시 활성화됩니다. 단 IAP 상품 등록은 Apple/Google 콘솔에서 직접 해야 하며, 앱 심사 통과 후에야 "
    "사용자가 실제 구매 가능."
)

heading("Q4. 한국 외부결제 (Toss를 앱 안에서) 가능한가?", level=3)
body(
    "Apple은 한국 한정 'External Purchase Entitlement' 정책 (2022~) 있음. "
    "Google도 'User Choice Billing' 한국 적용 중. 단, 신청·심사 절차 복잡하고 수수료 절감폭 크지 않음. "
    "1차 출시는 StoreKit/Play Billing 기본 권장. 외부결제는 Phase C로 분리."
)

heading("Q5. 토큰 환불 정책은?", level=3)
body(
    "Apple/Google: 사용자가 앱마켓에 직접 환불 요청 → 7일 내 자동 환불. INPICK 서버는 RTDN/ASN으로 "
    "환불 이벤트 수신 후 미사용 토큰만 자동 회수, 사용한 토큰은 reconciliation case로 관리자 수동 검토."
)

heading("Q6. 앱 심사 리젝되면?", level=3)
body(
    "iOS 심사 리젝 주요 사유: 단순 WebView 앱처럼 보임, 디지털 상품을 외부 결제로 판매, "
    "개인정보처리방침 부실. → 1차 출시는 앱 내 직접 결제 X, '웹에서 구매 후 앱에서 사용' 모델 권장."
)

doc.add_page_break()

# ═════════════════════════════════════════════════
# 8. 연락처 + 다음 단계
# ═════════════════════════════════════════════════
heading("8. 막히면 알려주세요", level=1)
body("각 단계마다 막히는 부분 있으면 알려주세요. 캡처/에러 메시지와 함께:")
bullet("Vercel 환경변수 등록 시 오류")
bullet("Apple Developer / Google Play 가입 절차 막힘")
bullet("App Store / Play Console 입력 양식 모르겠음")
bullet("결제 테스트 시 안 되는 경우")
bullet("심사 리젝 사유 분석")

body("준비 완료 상태:", bold=True)
bullet("코드/DB/API — 모두 작성 완료 ✓")
bullet("관리자 대시보드 (/admin/pricing, /admin/payment-center, /admin/members) — 구현 완료 ✓")
bullet("결제 검증 (Apple/Google 서버 API 실제 호출) — 코드 완료 ✓")
bullet("환경변수만 등록하면 즉시 활성화")

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run("문서 끝 — 2026-05-19 작성")
r2.italic = True
r2.font.size = Pt(9)

# 저장
out_path = r"C:\Users\user\Desktop\INPICK_결제_대표님_작업_가이드_20260519.docx"
doc.save(out_path)
print(f"[OK] 저장 완료: {out_path}")
print(f"[size] {os.path.getsize(out_path):,} bytes")
